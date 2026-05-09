from flask import Flask, request, jsonify, send_file
import os
import io
import requests
from bs4 import BeautifulSoup
from docx import Document
import google.generativeai as genai
import json
from flask_cors import CORS
import streamlit as st
from dotenv import load_dotenv


# Load .env file
load_dotenv()


app = Flask(__name__)
CORS(app)  # allow Streamlit (different port) to call this

# ---- Gemini setup ----
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

model = genai.GenerativeModel("models/gemma-4-26b-a4b-it")  # or gemini-1.5-pro


# ---- Helper: extract text from webpage ----
def extract_page_text(url: str) -> str:
    resp = requests.get(url, headers={"User-Agent": "Mozilla/5.0"})
    resp.raise_for_status()
    html = resp.text

    soup = BeautifulSoup(html, "html.parser")

    # SIMPLE VERSION: take text from main tags
    texts = []

    # try common content containers first (you can tune this later)
    candidates = soup.select("article, .content, .post-content, .entry-content")
    if candidates:
        for c in candidates:
            texts.append(c.get_text(separator="\n", strip=True))
    else:
        # fallback: all p + headings + list items
        for tag in soup.find_all(["h1", "h2", "h3", "p", "li"]):
            txt = tag.get_text(separator=" ", strip=True)
            if txt:
                texts.append(txt)

    page_text = "\n".join(texts)
    # optional: cut very large pages
    max_chars = 30000
    return page_text[:max_chars]


# ---- Route 1: scrape ----
@app.post("/scrape")
def scrape():
    data = request.get_json()
    url = data.get("url")
    if not url:
        return jsonify({"error": "url is required"}), 400

    try:
        page_text = extract_page_text(url)
        return jsonify({"pageText": page_text})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    
# Helper to parse LLM JSON output safely
# def parse_llm_json(raw_text: str):
#     import re, json

#     # Remove markdown fences
#     raw_text = re.sub(r"```json|```", "", raw_text).strip()

#     # Extract JSON array
#     match = re.search(r"\[\s*{.*}\s*\]", raw_text, re.DOTALL)
#     if not match:
#         raise ValueError("No JSON array found")

#     json_text = match.group()

#     # Normalize multiline strings
#     json_text = json_text.replace("\n", "\\n").replace("\t", "\\t")

#     return json.loads(json_text)

# ---- Route 2: answer questions with Gemini ----
@app.post("/answer-questions")
def answer_questions():

    try:
        data = request.get_json(force=True)

        page_text = str(data.get("pageText", ""))
        questions = data.get("questions", [])

        if not page_text or not questions:
            return jsonify({
                "error": "pageText and questions are required"
            }), 400

        qa_list = []

        for q in questions:

            prompt = f"""
            Answer this question briefly from the text.

            TEXT:
            {page_text[:3000]}

            QUESTION:
            {q}
            """

            response = model.generate_content(prompt)

            answer = response.text.strip()

            import re

            # Remove major reasoning junk
            patterns_to_remove = [
                r"Source text provided:.*",
                r"Input text:.*",
                r"Constraint:.*",
                r"Sentence \d+:.*",
                r"Draft \d+:.*",
                r"Brief\?.*",
                r"From the text\?.*",
                r"The text says:.*",
            ]

            for pattern in patterns_to_remove:
                answer = re.sub(pattern, "", answer, flags=re.IGNORECASE)

            cleaned_lines = []

            for line in answer.split("\n"):

                line = line.strip()

                if not line:
                    continue

                bad_starts = [
                    "Source",
                    "Input",
                    "Constraint",
                    "Sentence",
                    "Draft",
                    "Brief",
                    "Relevant",
                    "Irrelevant",
                    "The text says",
                    "Question:",
                    "TEXT:",
                ]

                skip = False

                for bad in bad_starts:
                    if line.startswith(bad):
                        skip = True
                        break

                # Skip quoted junk
                if line.startswith('"') and line.endswith('"'):
                    skip = True

                # Skip stars/bullets
                if line.startswith("*"):
                    skip = True

                if not skip:
                    cleaned_lines.append(line)

            answer = " ".join(cleaned_lines).strip()

            # Remove extra spaces
            answer = re.sub(r"\s+", " ", answer)

            # Keep only last proper sentence if repeated
            sentences = answer.split(".")

            if len(sentences) > 2:
                answer = ".".join(sentences[-2:]).strip()

            if answer and not answer.endswith("."):
                answer += "."

            # limit size
            answer = answer

            found = True

            if len(answer) < 15:
                found = False

            qa_list.append({
                "question": q,
                "answer": answer,
                "found": found
            })

        return jsonify({
            "qa": qa_list
        })

    except Exception as e:

        print("ERROR:", str(e))

        return jsonify({
            "error": "Failed to answer questions",
            "details": str(e)
        }), 500


# ---- Route 3: export DOCX ----
@app.post("/export-docx")
def export_docx():
    data = request.get_json()
    qa_list = data.get("qa", [])
    title = data.get("title", "Study Notes")

    doc = Document()
    doc.add_heading(title, level=1)

    for item in qa_list:
        q = item.get("question", "")
        a = item.get("answer", "")
        found = item.get("found", True)

        doc.add_heading(q, level=2)
        if not found:
            doc.add_paragraph("(Not clearly available in the given page.)")
        doc.add_paragraph(a)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return send_file(
        buf,
        as_attachment=True,
        download_name=f"{title.replace(' ', '_')}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
@app.route("/generate-questions", methods=["POST"])
def generate_questions():
    try:
        data = request.get_json(force=True)

        page_text = str(data.get("pageText", ""))
        num_questions = int(data.get("num_questions", 5))
        marks = int(data.get("marks", 10))
        difficulty = data.get("difficulty", "exam")

        if not page_text:
            return jsonify({"error": "pageText is required"}), 400

        # Smaller cleaner prompt for free-tier Gemma
        prompt = f"""
Generate {num_questions} university exam questions from the text below.

Rules:
- Questions only
- No answers
- No markdown
- One question per line
- Suitable for {marks} marks
- Difficulty level: {difficulty}

TEXT:
{page_text[:12000]}
"""

        response = model.generate_content(prompt)

        if not response.text:
            return jsonify({"error": "Empty response from model"}), 500

        raw_text = response.text.strip()

        print("RAW GEMINI QUESTIONS OUTPUT:\n", raw_text)

        # -------- Extract Questions Safely --------

        lines = raw_text.split("\n")

        questions = []

        for line in lines:
            line = line.strip()

            # Skip empty lines
            if not line:
                continue

            # Skip reasoning/checklist/debug lines
            if (
                line.startswith("*")
                or line.startswith("[")
                or line.startswith("]")
                or "Valid JSON" in line
                or "Pass" in line
                or "Constraint Check" in line
                or "Starts with" in line
                or "Ends with" in line
                or "Simple string" in line
                or "markdown" in line.lower()
                or "json" in line.lower()
            ):
                continue

            # Remove bullets / numbering
            line = line.lstrip("-•1234567890. ")

            # Remove quotes/comma
            cleaned = line.strip('",')

            # Keep only question-like lines
            if (
                len(cleaned) > 15
                and (
                    cleaned.startswith("Explain")
                    or cleaned.startswith("Describe")
                    or cleaned.startswith("Discuss")
                    or cleaned.startswith("Differentiate")
                    or cleaned.startswith("Compare")
                    or cleaned.startswith("What")
                    or cleaned.startswith("How")
                    or cleaned.startswith("Why")
                    or cleaned.startswith("Define")
                )
            ):
                questions.append(cleaned)

        # Remove duplicates
        questions = list(dict.fromkeys(questions))

        # Limit to requested number
        questions = questions[:num_questions]

        # Final fallback
        if not questions:
            return jsonify({
                "error": "Could not extract questions",
                "raw_output": raw_text
            }), 500

        return jsonify({"questions": questions})

    except Exception as e:
        print("ERROR generating questions:", str(e))

        return jsonify({
            "error": "Failed to generate questions",
            "details": str(e)
        }), 500


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
    # app.run(debug=True, port=5000)

